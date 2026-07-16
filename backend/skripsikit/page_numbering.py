from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from .constants import SECTPR_CHILD_ORDER_TAGS

SECTPR_CHILD_ORDER = [qn(f'w:{tag}') for tag in SECTPR_CHILD_ORDER_TAGS]


def insert_in_schema_order(sectPr, new_element):
    new_tag_order = SECTPR_CHILD_ORDER.index(new_element.tag)
    insert_pos = len(sectPr)
    for i, child in enumerate(sectPr):
        if child.tag in SECTPR_CHILD_ORDER:
            child_order = SECTPR_CHILD_ORDER.index(child.tag)
            if child_order > new_tag_order:
                insert_pos = i
                break
    sectPr.insert(insert_pos, new_element)


def set_page_number_format(section, fmt, restart=False):
    sectPr = section._sectPr
    for old in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(old)
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), fmt)
    if restart:
        pgNumType.set(qn('w:start'), '1')
    insert_in_schema_order(sectPr, pgNumType)


def add_page_number_field(paragraph, alignment, font_name='Times New Roman', font_size=12):
    """Isi paragraph header/footer dengan field code nomor halaman, font eksplisit."""
    paragraph.alignment = alignment
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)

    run = paragraph.add_run()
    run.font.name = font_name
    run.font.size = Pt(font_size)

    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)

    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    run._element.append(fldChar_begin)
    run._element.append(instrText)
    run._element.append(fldChar_end)


def clear_paragraph(paragraph):
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def setup_preliminary(section, restart=False):
    set_page_number_format(section, 'lowerRoman', restart=restart)

    section.footer.is_linked_to_previous = False
    footer_para = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    add_page_number_field(footer_para, WD_ALIGN_PARAGRAPH.CENTER)

    section.header.is_linked_to_previous = False
    header_para = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
    clear_paragraph(header_para)


def setup_chapter_section(section, is_first_chapter):
    set_page_number_format(section, 'decimal', restart=is_first_chapter)
    section.different_first_page_header_footer = True

    section.footer.is_linked_to_previous = False
    footer_para = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    clear_paragraph(footer_para)

    section.header.is_linked_to_previous = False
    header_para = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
    add_page_number_field(header_para, WD_ALIGN_PARAGRAPH.RIGHT)

    section.first_page_footer.is_linked_to_previous = False
    fp_footer_para = (section.first_page_footer.paragraphs[0]
                       if section.first_page_footer.paragraphs
                       else section.first_page_footer.add_paragraph())
    add_page_number_field(fp_footer_para, WD_ALIGN_PARAGRAPH.CENTER)

    section.first_page_header.is_linked_to_previous = False
    fp_header_para = (section.first_page_header.paragraphs[0]
                       if section.first_page_header.paragraphs
                       else section.first_page_header.add_paragraph())
    clear_paragraph(fp_header_para)


def configure_page_numbers(doc, section_points):
    """
    Petakan tiap heading BAB langsung ke section index yang memuatnya,
    dengan menelusuri dokumen paragraph-per-paragraph -- lebih robust
    dibanding menghitung lewat rumus posisi.
    """
    sections = doc.sections
    point_by_index = {p['index']: p for p in section_points}

    mapped = []
    current_section = 0
    for i, para in enumerate(doc.paragraphs):
        if i in point_by_index:
            mapped.append((point_by_index[i], current_section))
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
            current_section += 1

    if not mapped:
        print("✗ Tidak ada section yang berhasil terpetakan, proses dihentikan.")
        return

    first_chapter_section = mapped[0][1]

    for i in range(0, first_chapter_section):
        setup_preliminary(sections[i], restart=(i == 0))
    print(f"✓ Section 0-{first_chapter_section - 1} (Preliminary) → Roman, footer center")

    for pos, (point, idx) in enumerate(mapped):
        is_first = (pos == 0)
        setup_chapter_section(sections[idx], is_first_chapter=is_first)
        print(f"✓ {point['text']:30} → Arabic (section {idx})")